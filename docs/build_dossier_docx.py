# -*- coding: utf-8 -*-
"""Build a comprehensive Hazel Gym project dossier in DOCX format.

This document is intentionally more detailed than the final TFC report. It is a
single reference file for understanding the whole project, its decisions,
architecture, validation history and delivery state.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DATABASE = ROOT / "database"
BACKEND = ROOT / "backend"
FRONTEND = ROOT / "frontend"
MOBILE = ROOT / "mobile-android"
OUTPUT = DOCS / "HazelGym_Dossier_Maestro_Proyecto.docx"

ACCENT = RGBColor(255, 77, 46)
DARK = RGBColor(13, 13, 20)
MUTED = RGBColor(99, 111, 132)
GREEN = RGBColor(34, 204, 102)
BLUE = RGBColor(34, 102, 255)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.5)
    if color is not None:
        run.font.color.rgb = color


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for style_name, size, color in [
        ("Title", 25, DARK),
        ("Subtitle", 12, MUTED),
        ("Heading 1", 17, DARK),
        ("Heading 2", 13.5, DARK),
        ("Heading 3", 11.5, DARK),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        if "Heading" in style_name:
            style.font.bold = True
            style.paragraph_format.space_before = Pt(11)
            style.paragraph_format.space_after = Pt(4)

    header = section.header.paragraphs[0]
    header.text = "Hazel Gym · Dossier maestro del proyecto"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.add_run("Página ")
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = MUTED
    add_page_number(footer)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = DARK

    p = doc.add_paragraph()
    p.style = "Subtitle"
    p.add_run(subtitle)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_p(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic


def add_bullets(doc: Document, items: Iterable[str], level_style: str = "List Bullet") -> None:
    for item in items:
        p = doc.add_paragraph(style=level_style)
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_note(doc: Document, title: str, body: str, fill: str = "FFF1EC") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    r = p.add_run(title + ". ")
    r.bold = True
    r.font.color.rgb = ACCENT
    r.font.size = Pt(10)
    r2 = p.add_run(body)
    r2.font.size = Pt(10)


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.35)
    p.paragraph_format.right_indent = Cm(0.35)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(44, 52, 66)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_shading(hdr[i], "F2F4F7")
        set_cell_text(hdr[i], header, bold=True, color=DARK)
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            hdr[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if widths:
                cells[i].width = Inches(widths[i])


def add_image(doc: Document, image_path: Path, caption: str, width: float = 6.2) -> None:
    if not image_path.exists():
        return
    try:
        doc.add_picture(str(image_path), width=Inches(width))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        r.italic = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = MUTED
    except Exception as exc:
        add_note(doc, "Imagen no insertada", f"No se pudo insertar {image_path.name}: {exc}", fill="FFF6D7")


def list_files(folder: Path, pattern: str) -> list[str]:
    if not folder.exists():
        return []
    return [str(path.relative_to(ROOT)).replace("\\", "/") for path in sorted(folder.glob(pattern))]


def first_heading(markdown: Path) -> str:
    try:
        for line in markdown.read_text(encoding="utf-8").splitlines():
            line = line.strip()
            if line.startswith("#"):
                return line.lstrip("#").strip()
    except Exception:
        return markdown.stem
    return markdown.stem


def build_doc() -> None:
    doc = Document()
    configure_document(doc)

    add_title(
        doc,
        "Hazel Gym · Dossier maestro del proyecto",
        "Documento interno de comprensión total: arquitectura, decisiones, funcionalidades, despliegue, validación y material recomendado para el informe final.",
    )
    add_table(
        doc,
        ["Campo", "Valor"],
        [
            ["Proyecto", "Hazel Gym"],
            ["Autor", "Ander"],
            ["Fecha de generación", date.today().strftime("%d/%m/%Y")],
            ["Finalidad", "Base de estudio y redacción para el informe final del TFC/TFG. No sustituye a la memoria de entrega."],
            ["Estado", "Backend en AWS validado, frontend web desplegado, app Android conectada a API remota y APK de demostración generada."],
        ],
        widths=[1.8, 4.9],
    )

    add_heading(doc, "Cómo usar este documento", 1)
    add_p(
        doc,
        "Este dossier está escrito como si fuera el cuaderno técnico completo del proyecto. Su objetivo es que, al leerlo, puedas explicar con seguridad qué se ha construido, por qué se ha construido así, cómo se conectan sus partes y qué evidencias conviene mostrar en la defensa.",
    )
    add_bullets(
        doc,
        [
            "Usa las secciones 1 a 4 para redactar introducción, análisis, requisitos y diseño.",
            "Usa las secciones 5 a 8 para explicar base de datos, backend, app móvil y web.",
            "Usa las secciones 9 a 12 para explicar despliegue AWS, seguridad, pruebas y problemas resueltos.",
            "Usa las secciones 13 a 16 como banco de capturas, guion de defensa, anexos y mejoras futuras.",
        ]
    )

    add_heading(doc, "Índice orientativo", 1)
    add_numbered(
        doc,
        [
            "Resumen ejecutivo",
            "Contexto, motivación y objetivos",
            "Alcance funcional y roles",
            "Arquitectura general",
            "Base de datos y modelo de información",
            "Backend Spring Boot",
            "Aplicación móvil Android",
            "Aplicación web React",
            "Despliegue AWS e infraestructura",
            "Seguridad y configuración",
            "Validación funcional, pruebas e incidencias",
            "Decisiones técnicas y alternativas descartadas",
            "Estado de entrega y guía de demostración",
            "Material visual recomendado para la memoria",
            "Preguntas de defensa y respuestas preparadas",
            "Anexos e inventario de fuentes",
        ]
    )

    add_heading(doc, "Equivalencia con la estructura recomendada del TFC", 1)
    add_p(
        doc,
        "La plantilla de referencia del TFC pide una memoria ordenada por introducción, análisis/diseño, base de datos, backend, interfaces, despliegue, seguridad, pruebas, conclusiones y anexos. Este dossier usa esa misma lógica, pero amplía cada bloque con decisiones, errores solucionados y recomendaciones prácticas para que luego puedas seleccionar qué llevar a la memoria final.",
    )
    add_table(
        doc,
        ["Bloque de la memoria final", "Dónde encontrarlo en este dossier"],
        [
            ["Introducción, motivación y objetivos", "Secciones 1 y 2."],
            ["Análisis, roles, requisitos y diseño", "Secciones 3 y 4."],
            ["Base de datos", "Sección 5."],
            ["Backend", "Sección 6."],
            ["App móvil", "Sección 7."],
            ["Web", "Sección 8."],
            ["Despliegue e infraestructura", "Sección 9."],
            ["Seguridad", "Sección 10."],
            ["Pruebas y validación", "Sección 11."],
            ["Conclusiones, defensa y anexos", "Secciones 12 a 16."],
        ],
        widths=[2.7, 4.0],
    )

    add_heading(doc, "1. Resumen ejecutivo", 1)
    add_p(
        doc,
        "Hazel Gym es una plataforma de gestión para un gimnasio que integra backend REST, aplicación móvil Android, aplicación web y despliegue en AWS. El sistema cubre tres perfiles principales: cliente, entrenador y administrador. La idea central es digitalizar operaciones habituales de un gimnasio: autenticación, gestión de usuarios, máquinas, clases, rutinas, cuotas, códigos QR y asistencias.",
    )
    add_p(
        doc,
        "El proyecto no se limita a una maqueta: cuenta con backend en Spring Boot, base de datos MySQL en Amazon RDS, API desplegada en Elastic Beanstalk, API Gateway como puerta HTTPS, frontend web desplegado en AWS Amplify y app Android generable como APK. Esta composición permite mostrar el proyecto desde navegador y desde móvil sin depender únicamente de localhost.",
    )
    add_note(
        doc,
        "Idea clave",
        "La app móvil está pensada para el uso dentro del gimnasio, especialmente el escaneo de QR. La web está pensada para escritorio y gestión. Ambas consumen la misma API REST, lo que mantiene coherencia funcional y evita duplicar lógica de negocio.",
    )

    add_heading(doc, "2. Contexto, motivación y objetivos", 1)
    add_heading(doc, "2.1 Contexto", 2)
    add_p(
        doc,
        "El proyecto nace como trabajo final del ciclo DAM, con el objetivo de construir una solución completa y defendible que combine tecnologías vistas en el ciclo con herramientas profesionales actuales. El contexto funcional es un gimnasio que necesita centralizar su gestión diaria y ofrecer a clientes y entrenadores una experiencia digital sencilla.",
    )
    add_heading(doc, "2.2 Motivación", 2)
    add_bullets(
        doc,
        [
            "Crear una aplicación con utilidad real y no solo una demostración aislada.",
            "Practicar un flujo completo de desarrollo: análisis, diseño, base de datos, backend, frontend móvil, frontend web, despliegue, pruebas y documentación.",
            "Aplicar tecnologías propias del ciclo DAM, especialmente Java, SQL, Kotlin y Android.",
            "Aprender a llevar el proyecto a la nube con AWS y dejarlo accesible para una demo pública.",
        ],
    )
    add_heading(doc, "2.3 Objetivos", 2)
    add_table(
        doc,
        ["Objetivo", "Cómo se cumple"],
        [
            ["Gestión de usuarios y roles", "Modelo de usuarios con roles CLIENT, TRAINER y ADMIN, autenticación JWT y paneles específicos."],
            ["Gestión del gimnasio", "CRUD de máquinas, clases, sesiones, rutinas, asignaciones, cuotas y asistencias."],
            ["Uso de códigos QR", "QR de entrada, QR de máquina y QR de sesión/clase, con lectura desde app móvil."],
            ["Aplicación móvil nativa", "Android con Kotlin, Jetpack Compose, Retrofit, DataStore, CameraX y ML Kit."],
            ["Aplicación web", "React/Vite con paneles por rol y consumo de API real."],
            ["Despliegue real", "RDS, Elastic Beanstalk, API Gateway y Amplify en AWS."],
            ["Validación", "Smoke test completo remoto, builds de frontend/backend y APK generada."],
        ],
        widths=[2.0, 4.7],
    )
    add_heading(doc, "2.4 Línea temporal resumida", 2)
    add_table(
        doc,
        ["Fecha / fase", "Avance principal"],
        [
            ["30 enero", "Definición inicial del proyecto Hazel Gym, roles principales y casos de uso base."],
            ["Abril", "Refinamiento de clases, sesiones, QR y alcance funcional realista."],
            ["3 mayo", "Modelo entidad-relación con 10 entidades y descarte de gamificación/puntos para mantener foco MVP."],
            ["4 mayo", "Preparación de MySQL local, scripts SQL y datos seed."],
            ["6-7 mayo", "Backend Spring Boot limpio: entidades, repositorios, servicios, controladores, seguridad y Swagger."],
            ["10-11 mayo", "Corrección de JWT, tipos de columnas, Swagger, smoke tests y conexión estable local."],
            ["12-13 mayo", "App Android: login, roles, navegación, QR, gestión de usuarios, clases, rutinas y asignaciones."],
            ["17-20 mayo", "Frontend web React/Vite/pnpm y documentación de React para defensa."],
            ["21-25 mayo", "Despliegue AWS: RDS, Elastic Beanstalk, API Gateway, Amplify y CORS."],
            ["25-26 mayo", "Validación final remota, APK conectada a AWS, datos demo y preparación de informe."],
        ],
        widths=[1.7, 5.0],
    )

    add_heading(doc, "3. Alcance funcional y roles", 1)
    add_heading(doc, "3.1 Roles del sistema", 2)
    add_table(
        doc,
        ["Rol", "Responsabilidad", "Funciones principales"],
        [
            ["CLIENT", "Usuario socio del gimnasio.", "Login, consulta de máquinas, rutinas asignadas, historial de asistencia, escaneo de QR de entrada/sesión/máquina, perfil."],
            ["TRAINER", "Entrenador que gestiona seguimiento de clientes.", "Consulta de clientes, sesiones, clases, rutinas, asignación de rutinas, seguimiento y perfil."],
            ["ADMIN", "Administrador del gimnasio.", "Gestión de usuarios, máquinas, clases, sesiones, QR, cuotas, asistencias y actividad del sistema."],
        ],
        widths=[1.2, 2.0, 3.5],
    )
    add_heading(doc, "3.2 Flujo funcional principal", 2)
    add_numbered(
        doc,
        [
            "El usuario accede con email y contraseña.",
            "El backend valida credenciales, genera JWT y devuelve los datos del usuario autenticado.",
            "La aplicación detecta el rol real del usuario desde el backend; no se confía en una selección visual del rol.",
            "La interfaz carga el panel correspondiente: cliente, entrenador o administrador.",
            "Las operaciones protegidas se ejecutan con el token JWT en la cabecera Authorization.",
            "Los códigos QR se gestionan desde administración y se leen principalmente desde la app móvil.",
        ],
    )
    add_heading(doc, "3.3 Funcionalidades por bloque", 2)
    add_table(
        doc,
        ["Bloque", "Funcionalidades"],
        [
            ["Autenticación", "Registro, login, consulta de usuario actual, validación de rol, cierre de sesión."],
            ["Usuarios", "Listado, creación, edición, eliminación y filtrado por rol según permisos."],
            ["Máquinas", "Catálogo, estado, grupo muscular, instrucciones, advertencias y recurso multimedia asociado."],
            ["Clases y sesiones", "Clases generales y sesiones programadas para fecha/hora concreta."],
            ["Rutinas", "Creación de rutinas, gestión de ejercicios y asignación a clientes."],
            ["QR", "Generación de QR de entrada, máquina y sesión; lectura desde móvil; modo manual de prueba en emulador."],
            ["Asistencias", "Registro de asistencia por QR y consulta de historial propio o global según rol."],
            ["Cuotas", "Gestión de cuotas/membresías desde administración y consulta por usuarios autorizados."],
        ],
        widths=[1.7, 5.0],
    )

    add_heading(doc, "4. Arquitectura general", 1)
    add_p(
        doc,
        "La arquitectura se ha dividido en capas y clientes independientes. El backend concentra la lógica de negocio y expone una API REST. La app móvil y la app web son clientes que consumen esa API. La base de datos está externalizada en Amazon RDS. Para la web desplegada en HTTPS se añadió API Gateway como proxy HTTPS hacia Elastic Beanstalk.",
    )
    add_code(
        doc,
        "Usuario web -> AWS Amplify (React/Vite HTTPS) -> /api-proxy -> API Gateway HTTPS -> Elastic Beanstalk HTTP -> Spring Boot -> RDS MySQL\n"
        "Usuario móvil -> APK Android -> API Gateway HTTPS -> Elastic Beanstalk HTTP -> Spring Boot -> RDS MySQL\n"
        "Desarrollador -> GitHub -> GitHub Actions / Amplify / Elastic Beanstalk -> artefactos desplegados",
    )
    add_image(doc, ROOT / "DiagramaGimnasio.drawio.png", "Diagrama general del gimnasio y flujo conceptual del sistema.")
    add_heading(doc, "4.1 Conexiones actuales", 2)
    add_table(
        doc,
        ["Elemento", "URL / Ubicación", "Uso"],
        [
            ["Frontend web", "https://main.d1mithns8dqv1b.amplifyapp.com/", "Cliente web desplegado en Amplify."],
            ["API Gateway", "https://k7edn14r3k.execute-api.eu-west-1.amazonaws.com", "Endpoint HTTPS común para web y Android."],
            ["Elastic Beanstalk", "http://hazelgym-backend.eu-west-1.elasticbeanstalk.com", "Backend Spring Boot ejecutado en AWS."],
            ["RDS MySQL", "hazelgym-db.cpsq2kmkisyt.eu-west-1.rds.amazonaws.com:3306", "Base de datos remota hazelgym."],
            ["Android", "mobile-android/app/build/outputs/apk/debug", "APK de demo generada y conectada a API Gateway."],
        ],
        widths=[1.6, 3.2, 1.9],
    )
    add_note(
        doc,
        "Dominio hazelgym.app",
        "El dominio existe y actualmente se ha dejado como posible mejora. Un reenvío URL puede redirigir al usuario a Amplify, pero para que el dominio sea nativo conviene añadirlo como custom domain en Amplify, configurar DNS y añadir https://hazelgym.app al CORS del backend.",
        fill="EEF6FF",
    )

    add_heading(doc, "5. Base de datos y modelo de información", 1)
    add_p(
        doc,
        "La base de datos utiliza MySQL porque el dominio del problema es claramente relacional: usuarios, roles, máquinas, clases, sesiones, asistencias, rutinas, cuotas y QR tienen relaciones bien definidas. Además, MySQL encaja con Amazon RDS y con lo trabajado durante el ciclo.",
    )
    add_image(doc, ROOT / "Diagrama_entidad-relacion_imagen.png", "Diagrama entidad-relación utilizado como referencia del modelo de datos.")
    add_heading(doc, "5.1 Entidades principales", 2)
    add_table(
        doc,
        ["Entidad", "Descripción", "Relaciones principales"],
        [
            ["roles", "Catálogo de roles del sistema.", "Un rol puede estar asociado a muchos usuarios."],
            ["usuarios", "Personas registradas con email, contraseña cifrada, nombre y rol.", "Pertenece a un rol; genera asistencias; puede recibir rutinas."],
            ["maquinas", "Máquinas del gimnasio con estado, grupo muscular e instrucciones.", "Puede tener QR de tipo máquina."],
            ["clases", "Actividades generales ofrecidas por el gimnasio.", "Tiene sesiones programadas."],
            ["class_sessions", "Sesiones concretas de una clase en fecha/hora.", "Puede tener QR de sesión y asistencias."],
            ["qr_codes", "Códigos QR con tipo ENTRY, MACHINE o CLASS_SESSION.", "Apunta opcionalmente a máquina o sesión."],
            ["asistencias", "Registro de entrada o asistencia generado al usar QR.", "Relaciona usuario y QR."],
            ["rutinas", "Rutinas de entrenamiento creadas por admin/entrenador.", "Se asignan a clientes."],
            ["routine_assignments", "Asignación de rutina a cliente.", "Relaciona usuario cliente y rutina."],
            ["membership_fees", "Cuotas o planes de membresía.", "Consultables y gestionables desde la app."],
        ],
        widths=[1.5, 2.8, 2.4],
    )
    add_heading(doc, "5.2 Scripts SQL", 2)
    sql_files = list_files(DATABASE, "*.sql")
    add_table(
        doc,
        ["Script", "Función"],
        [[f, {
            "01_create_database.sql": "Crea la base de datos hazelgym.",
            "02_schema.sql": "Define tablas, claves primarias y relaciones.",
            "03_seed.sql": "Carga roles, usuarios y datos iniciales.",
            "04_prepare_existing_schema.sql": "Ajustes de compatibilidad en esquemas existentes.",
            "05_demo_machine_media.sql": "Añade instrucciones, advertencias y recursos de máquinas.",
            "06_cleanup_smoke_data.sql": "Limpia datos generados por pruebas smoke.",
            "07_prepare_delivery_demo_data.sql": "Deja datos presentables para demo final.",
        }.get(Path(f).name, "Script de apoyo del proyecto.")] for f in sql_files],
        widths=[2.5, 4.2],
    )
    add_note(
        doc,
        "Incidencia SQL Safe Updates",
        "MySQL Workbench bloqueó algunos DELETE/UPDATE con el error 1175 porque el modo seguro exige condiciones basadas en clave. Se resolvió usando condiciones con id o desactivando temporalmente SQL_SAFE_UPDATES de forma controlada.",
        fill="FFF6D7",
    )

    add_heading(doc, "6. Backend Spring Boot", 1)
    add_p(
        doc,
        "El backend es el núcleo del proyecto. Se ha implementado con Java 17 y Spring Boot 4.0.6 siguiendo una estructura típica por capas: controladores REST, servicios, repositorios JPA, entidades y configuración de seguridad. Esta elección permite una API mantenible, documentada y fácil de desplegar como JAR.",
    )
    add_heading(doc, "6.1 Tecnologías backend", 2)
    add_table(
        doc,
        ["Tecnología", "Uso", "Motivo"],
        [
            ["Java 17", "Lenguaje del backend.", "Versión moderna, estable y compatible con Elastic Beanstalk Java 17."],
            ["Spring Boot", "Aplicación REST y configuración automática.", "Reduce boilerplate y facilita seguridad, JPA y despliegue."],
            ["Spring Data JPA", "Acceso a base de datos.", "Repositorios simples y entidades relacionales."],
            ["Spring Security", "Protección por JWT y roles.", "Control de acceso robusto y estándar."],
            ["jjwt", "Generación y validación de tokens JWT.", "Token stateless compartido por web y móvil."],
            ["Springdoc/OpenAPI", "Swagger y documentación API.", "Permite probar endpoints desde navegador."],
            ["Maven", "Gestión de dependencias y build.", "Ecosistema Java conocido y compatible con CI."],
            ["MySQL Connector/J", "Conexión a MySQL/RDS.", "Driver oficial para la base de datos elegida."],
        ],
        widths=[1.7, 2.3, 2.7],
    )
    add_heading(doc, "6.2 Configuración por variables de entorno", 2)
    add_p(
        doc,
        "El backend no fija credenciales sensibles en código. Lee la conexión a base de datos, usuario, contraseña, JWT y CORS desde variables de entorno. Esto permite usar la misma aplicación en local, Elastic Beanstalk y pruebas sin cambiar el código fuente.",
    )
    add_code(
        doc,
        "MYSQL_URL=jdbc:mysql://.../hazelgym?useSSL=true&serverTimezone=UTC&characterEncoding=utf8\n"
        "MYSQL_USERNAME=admin_hazelgym\n"
        "MYSQL_PASSWORD=********\n"
        "JWT_SECRET=********\n"
        "JWT_EXPIRATION_MS=86400000\n"
        "APP_CORS_ALLOWED_ORIGINS=http://localhost:5173,https://main.d1mithns8dqv1b.amplifyapp.com",
    )
    add_heading(doc, "6.3 API REST", 2)
    add_table(
        doc,
        ["Recurso", "Endpoint base", "Descripción"],
        [
            ["Autenticación", "/api/auth", "Login, registro y usuario actual."],
            ["Usuarios", "/api/users", "Gestión y consulta de usuarios por rol."],
            ["Máquinas", "/api/machines", "Catálogo y gestión de máquinas."],
            ["Clases", "/api/classes", "Gestión de clases generales."],
            ["Sesiones", "/api/class-sessions", "Sesiones programadas de clases."],
            ["QR", "/api/qr-codes", "Creación, consulta y eliminación de códigos QR."],
            ["Asistencias", "/api/attendances", "Registro y consulta de asistencias."],
            ["Rutinas", "/api/routines", "Gestión de rutinas."],
            ["Asignaciones", "/api/routine-assignments", "Asignación de rutinas a clientes."],
            ["Cuotas", "/api/membership-fees", "Gestión de cuotas/membresías."],
        ],
        widths=[1.4, 2.3, 3.0],
    )
    add_heading(doc, "6.4 Seguridad backend", 2)
    add_bullets(
        doc,
        [
            "Las contraseñas se almacenan con BCrypt, nunca en texto plano.",
            "El login devuelve un JWT firmado; el cliente lo envía en Authorization: Bearer <token>.",
            "Los endpoints públicos se limitan a login, registro, Swagger/OpenAPI y la raíz informativa.",
            "El resto de rutas exige autenticación y, en muchos casos, rol adecuado.",
            "La raíz / inicialmente devolvía 403 en Elastic Beanstalk porque estaba protegida; se añadió un RootController y se permitió públicamente para mostrar el estado de la API.",
        ],
    )
    add_note(
        doc,
        "Error JWT resuelto",
        "Al principio se usó una clave demasiado corta y jjwt lanzó WeakKeyException. Se corrigió usando un secreto de longitud suficiente para HMAC-SHA, leído desde variable de entorno.",
        fill="FFF6D7",
    )

    add_heading(doc, "7. Aplicación móvil Android", 1)
    add_p(
        doc,
        "La app móvil se ha desarrollado en Kotlin con Jetpack Compose porque el proyecto pertenece a DAM y esta tecnología ya estaba dentro del aprendizaje del autor. Es la pieza más natural para funciones de gimnasio físico: escanear QR, consultar máquinas, registrar asistencia y acceder al perfil desde el teléfono.",
    )
    add_heading(doc, "7.1 Tecnologías Android", 2)
    add_table(
        doc,
        ["Tecnología", "Uso", "Por qué"],
        [
            ["Kotlin", "Lenguaje nativo Android.", "Moderno, expresivo y trabajado en DAM."],
            ["Jetpack Compose", "Interfaz declarativa.", "Permite pantallas rápidas, mantenibles y cercanas a Figma."],
            ["Navigation Compose", "Rutas internas.", "Navegación por pantallas y roles sin Activitys múltiples."],
            ["Retrofit + OkHttp", "Cliente HTTP.", "Integración clara con API REST."],
            ["DataStore", "Persistencia de sesión.", "Alternativa moderna y segura a SharedPreferences."],
            ["CameraX + ML Kit", "Escaneo QR.", "Permite lectura de códigos desde cámara real."],
            ["ZXing", "Generación/representación QR.", "Apoyo para mostrar códigos QR."],
        ],
        widths=[1.8, 2.3, 2.6],
    )
    add_heading(doc, "7.2 Estructura funcional móvil", 2)
    add_table(
        doc,
        ["Rol", "Pantallas/funciones"],
        [
            ["Cliente", "Login, inicio, historial de asistencias, escaneo QR, máquinas, instrucciones/recurso multimedia, rutinas asignadas y perfil."],
            ["Entrenador", "Inicio, clientes, clases/sesiones, rutinas, asignaciones de rutinas y perfil."],
            ["Administrador", "Inicio, usuarios, máquinas, clases, sesiones, QR, actividad, cuotas y perfil."],
        ],
        widths=[1.3, 5.4],
    )
    add_heading(doc, "7.3 Conexión remota móvil", 2)
    add_p(
        doc,
        "La app móvil dejó de depender de 10.0.2.2, que solo sirve para conectar el emulador Android con localhost del ordenador. La URL base de producción apunta a API Gateway, de forma que una APK instalada en un móvil real puede comunicarse con el backend remoto.",
    )
    add_code(doc, 'BuildConfig.API_BASE_URL = "https://k7edn14r3k.execute-api.eu-west-1.amazonaws.com/"')
    add_note(
        doc,
        "Login móvil",
        "Se eliminó el selector obligatorio de rol antes de entrar. Igual que en la web, el rol se detecta desde el backend tras autenticar. También se eliminó el texto 'Olvidaste tu contraseña?' porque el proyecto no implementa recuperación de credenciales.",
    )
    add_heading(doc, "7.4 APK", 2)
    add_p(
        doc,
        "La APK de demostración se generó en mobile-android/app/build/outputs/apk/debug. Se configuró el nombre de salida para que no quedara como app-debug.apk, sino como una APK identificable de Hazel Gym.",
    )

    add_heading(doc, "8. Aplicación web React", 1)
    add_p(
        doc,
        "La web se creó con React, Vite y TypeScript. Aunque Flutter Web se valoró inicialmente, React se eligió para mantener una base sencilla, fácil de desplegar en Amplify y más directa para una aplicación de escritorio. Se ha intentado que el código sea entendible para alguien que no había trabajado previamente con React.",
    )
    add_heading(doc, "8.1 Qué es React en este proyecto", 2)
    add_p(
        doc,
        "React divide la interfaz en componentes reutilizables. En Hazel Gym hay componentes de login, layout de panel, tarjetas, listados y paneles por rol. Cada componente recibe datos, muestra una parte de pantalla y reacciona a cambios de estado como login correcto, carga de datos o selección de una sección.",
    )
    add_table(
        doc,
        ["Concepto React", "Aplicación en Hazel Gym"],
        [
            ["Componente", "Bloques visuales como LoginScreen, AdminDashboard, ClientDashboard o TrainerDashboard."],
            ["Estado", "Usuario autenticado, token, sección activa, listados cargados y errores."],
            ["Props", "Datos que se pasan entre componentes: usuario, función logout, sección activa."],
            ["Servicios", "Archivos que hacen fetch a la API y aíslan la comunicación HTTP."],
            ["Build Vite", "Compila la web en frontend/dist para despliegue en Amplify."],
        ],
        widths=[1.8, 4.9],
    )
    add_heading(doc, "8.2 Configuración de API web", 2)
    add_p(
        doc,
        "La web usa VITE_API_BASE_URL. En local puede apuntar a localhost o API Gateway. En producción se usa /api-proxy, que Amplify redirige hacia API Gateway. Esto evita que el navegador bloquee peticiones por contenido mixto HTTPS/HTTP.",
    )
    add_code(
        doc,
        "VITE_API_BASE_URL=/api-proxy\n"
        "Amplify rewrite: /api-proxy/<*> -> https://k7edn14r3k.execute-api.eu-west-1.amazonaws.com/<*>",
    )
    add_heading(doc, "8.3 Funcionalidad web por rol", 2)
    add_table(
        doc,
        ["Rol", "Uso principal en web"],
        [
            ["Cliente", "Consulta de inicio, asistencia, máquinas, rutinas y perfil. El escaneo QR se considera más propio de móvil."],
            ["Entrenador", "Gestión/consulta de clientes, sesiones, rutinas y perfil."],
            ["Administrador", "Gestión completa de usuarios, máquinas, clases, cuotas, QR y actividad."],
        ],
        widths=[1.3, 5.4],
    )
    add_heading(doc, "8.4 Diseño web", 2)
    add_bullets(
        doc,
        [
            "Paleta coherente con la app móvil: naranja principal, fondos oscuros, azul para entrenador y verde para administrador.",
            "Pantallas responsive con layout de panel completo, evitando bloques que parezcan flotantes sobre fondo vacío.",
            "Logo añadido como favicon y dentro de las pantallas para reforzar identidad visual.",
            "Navegación lateral por secciones para mantener claridad en escritorio.",
        ],
    )

    add_heading(doc, "9. Despliegue AWS e infraestructura", 1)
    add_p(
        doc,
        "El despliegue se ha realizado en AWS para que el proyecto pueda mostrarse sin depender del ordenador local. La arquitectura final usa varios servicios, cada uno con una responsabilidad clara.",
    )
    add_table(
        doc,
        ["Servicio AWS", "Responsabilidad", "Motivo de elección"],
        [
            ["Amazon RDS MySQL", "Alojar la base de datos hazelgym.", "MySQL gestionado, compatible con el backend y fácil de conectar desde Spring."],
            ["Elastic Beanstalk", "Ejecutar el JAR de Spring Boot.", "Despliegue Java relativamente simple sin gestionar manualmente toda la infraestructura EC2."],
            ["API Gateway", "Ofrecer endpoint HTTPS estable hacia el backend HTTP.", "Evita mixed content en web y da URL limpia para Android."],
            ["AWS Amplify", "Hosting del frontend React.", "Integración con GitHub, builds automáticos y HTTPS."],
            ["AWS Budget", "Control de gasto.", "Evita sorpresas económicas durante la demo."],
        ],
        widths=[1.8, 2.5, 2.4],
    )
    add_heading(doc, "9.1 Por qué se añadió API Gateway", 2)
    add_p(
        doc,
        "Amplify sirve la web por HTTPS. Elastic Beanstalk expone el backend por HTTP. Los navegadores bloquean peticiones fetch desde una web HTTPS hacia un backend HTTP por mixed content. Además, Amplify no permite usar un destino HTTP directo en sus reglas de reverse proxy. La solución práctica fue insertar API Gateway, que ofrece HTTPS públicamente y reenvía internamente a Elastic Beanstalk HTTP.",
    )
    add_code(
        doc,
        "Amplify HTTPS (/api-proxy) -> API Gateway HTTPS -> Elastic Beanstalk HTTP -> Spring Boot -> RDS",
    )
    add_note(
        doc,
        "Incidencia API Gateway",
        "El login web devolvía 'Request method POST is not supported' porque la integración de API Gateway no incluía /{proxy}. Se corrigió dejando la ruta ANY /{proxy+} y la integración hacia http://hazelgym-backend.eu-west-1.elasticbeanstalk.com/{proxy}.",
        fill="FFF6D7",
    )
    add_heading(doc, "9.2 Datos de infraestructura", 2)
    add_table(
        doc,
        ["Elemento", "Valor"],
        [
            ["Región", "eu-west-1 (Irlanda)"],
            ["RDS endpoint", "hazelgym-db.cpsq2kmkisyt.eu-west-1.rds.amazonaws.com"],
            ["Base de datos", "hazelgym"],
            ["Usuario RDS", "admin_hazelgym"],
            ["Elastic Beanstalk", "http://hazelgym-backend.eu-west-1.elasticbeanstalk.com"],
            ["API Gateway", "https://k7edn14r3k.execute-api.eu-west-1.amazonaws.com"],
            ["Amplify", "https://main.d1mithns8dqv1b.amplifyapp.com/"],
            ["Dominio reservado", "hazelgym.app"],
            ["Budget", "hazelgym-budget, 30 USD/mes, alertas 80% y 100%"],
        ],
        widths=[1.8, 4.9],
    )
    add_heading(doc, "9.3 CI/CD y automatización", 2)
    add_p(
        doc,
        "El repositorio contiene workflows de GitHub Actions para integración continua, compilación de APK y despliegues. Amplify también queda conectado al repositorio para reconstruir automáticamente la web cuando se suben cambios a la rama configurada.",
    )
    add_table(
        doc,
        ["Workflow", "Propósito"],
        [
            [".github/workflows/ci.yml", "Validar backend/frontend en CI."],
            [".github/workflows/android-apk.yml", "Compilar APK Android usando API base remota."],
            [".github/workflows/deploy-backend-eb.yml", "Preparar despliegue del backend en Elastic Beanstalk."],
            [".github/workflows/deploy-frontend-s3.yml", "Alternativa documentada para hosting frontend en S3."],
        ],
        widths=[2.7, 4.0],
    )

    add_heading(doc, "10. Seguridad y configuración", 1)
    add_bullets(
        doc,
        [
            "JWT stateless para que web y móvil usen el mismo mecanismo de autenticación.",
            "Contraseñas cifradas con BCrypt.",
            "CORS controlado por variable de entorno para permitir solo orígenes conocidos.",
            "Credenciales reales fuera del repositorio mediante variables de entorno en local y AWS.",
            "Security Group de RDS configurado para permitir acceso desde Elastic Beanstalk y desde IP de trabajo durante desarrollo.",
            "API Gateway usado como capa HTTPS, aunque no como authorizer: la autorización real se mantiene en Spring Security.",
        ],
    )
    add_note(
        doc,
        "No duplicar seguridad",
        "No se añadió authorizer en API Gateway porque el backend ya valida sus propios JWT. Añadirlo antes de la entrega habría duplicado complejidad y requeriría adaptar tokens a OIDC/Cognito.",
        fill="EEF6FF",
    )

    add_heading(doc, "11. Validación funcional, pruebas e incidencias", 1)
    add_heading(doc, "11.1 Validación final remota", 2)
    add_p(
        doc,
        "La prueba smoke remota contra API Gateway pasó correctamente. Esto valida que la API pública HTTPS, Elastic Beanstalk, RDS, JWT, roles y operaciones principales están funcionando en conjunto.",
    )
    add_code(
        doc,
        'Smoke test contra https://k7edn14r3k.execute-api.eu-west-1.amazonaws.com\n'
        "[OK] OpenAPI docs are available\n"
        "[OK] Client registration returned token\n"
        "[OK] Client token authenticates /api/auth/me\n"
        "[OK] Client can list machines (3)\n"
        "[OK] Client can register attendance id=9\n"
        "[OK] Client users endpoint only returns own user\n"
        "[OK] Admin login returned token\n"
        "[OK] Admin can list users (4)\n"
        "[OK] Admin can create/delete machine, membership fee, class, session, QR, routine and assignment\n"
        "[OK] Client can register class session attendance\n"
        "Smoke test completed successfully",
    )
    add_heading(doc, "11.2 Incidencias importantes resueltas", 2)
    add_table(
        doc,
        ["Incidencia", "Causa", "Solución"],
        [
            ["JWT WeakKeyException", "Clave HMAC demasiado corta.", "Usar JWT_SECRET suficientemente largo."],
            ["MySQL Access denied", "Variable MYSQL_PASSWORD no aplicada o terminal distinta.", "Configurar variables de entorno en la sesión correcta."],
            ["Hibernate int/bigint", "Diferencia entre schema SQL y entidades Java.", "Alinear tipos de ID y scripts."],
            ["Swagger 500", "Incompatibilidad de Springdoc con versión de Spring.", "Ajustar dependencia compatible."],
            ["Android cleartext 10.0.2.2", "El emulador usaba HTTP local y política de red.", "Configurar acceso y después migrar a API Gateway HTTPS."],
            ["Gradle/OkHttp no resuelto", "Dependencia mal nombrada/caché.", "Corregir coordenadas y sincronizar."],
            ["Errores Compose", "Imports, smart casts y llamadas composable fuera de contexto.", "Refactorizar pantallas y ViewModels."],
            ["Git filename too long", "node_modules/build generados intentaban entrar al commit.", "Excluir artefactos generados y limpiar seguimiento."],
            ["Mixed content web", "Amplify HTTPS llamaba a EB HTTP.", "API Gateway como proxy HTTPS."],
            ["POST not supported en login", "Integración API Gateway sin /{proxy}.", "Actualizar integración a EB con /{proxy}."],
            ["Workbench error 1175", "Safe Updates activo.", "Usar condición por clave o SQL_SAFE_UPDATES controlado."],
        ],
        widths=[1.8, 2.4, 2.5],
    )

    add_heading(doc, "12. Decisiones técnicas y alternativas descartadas", 1)
    add_table(
        doc,
        ["Decisión", "Alternativa", "Motivo"],
        [
            ["Spring Boot", "Backend en Node/Express", "Mayor relación con DAM, seguridad/JPA integradas y despliegue Java directo."],
            ["MySQL/RDS", "SQLite o MongoDB", "Modelo relacional, integridad y compatibilidad con AWS."],
            ["JWT", "Sesiones servidor", "Más adecuado para web + móvil y backend stateless."],
            ["Jetpack Compose", "XML Android clásico", "Más moderno, productivo y trabajado en DAM."],
            ["React web", "Flutter Web", "Menor complejidad para web de gestión y despliegue Amplify."],
            ["pnpm", "npm", "Preferencia por seguridad, lockfile estricto y velocidad."],
            ["API Gateway", "Configurar HTTPS completo en EB", "Más rápido para la entrega y resuelve mixed content sin dominio/certificado inmediato."],
            ["QR en móvil", "QR completo en web", "La cámara/uso físico tiene más sentido en teléfono que en PC de escritorio."],
            ["Sin puntos/ranking", "Gamificación completa", "Se priorizó MVP funcional y defendible antes de ampliar alcance."],
        ],
        widths=[1.8, 1.8, 3.1],
    )

    add_heading(doc, "13. Estado de entrega y guía de demostración", 1)
    add_heading(doc, "13.1 Estado actual", 2)
    add_bullets(
        doc,
        [
            "Backend desplegado en AWS y validado con smoke test remoto.",
            "Base de datos RDS creada, cargada y conectada al backend.",
            "Frontend web desplegado en Amplify y conectado mediante /api-proxy.",
            "API Gateway operativo como endpoint HTTPS común.",
            "App Android conectada a API Gateway y APK generada.",
            "Datos demo preparados con usuarios y máquinas presentables.",
        ],
    )
    add_heading(doc, "13.2 Demo recomendada", 2)
    add_numbered(
        doc,
        [
            "Abrir la web en Amplify y hacer login como administrador.",
            "Mostrar dashboard, usuarios, máquinas, QR y actividad.",
            "Abrir Swagger o /api-docs para demostrar API documentada.",
            "Ejecutar o enseñar resultado del smoke test remoto.",
            "Instalar/abrir APK y hacer login desde móvil real.",
            "Mostrar catálogo de máquinas y detalle con instrucciones/recurso.",
            "Escanear o simular QR de máquina/sesión/entrada para demostrar flujo móvil.",
            "Cerrar explicando la arquitectura AWS: Amplify, API Gateway, Elastic Beanstalk y RDS.",
        ],
    )

    add_heading(doc, "14. Material visual recomendado para la memoria final", 1)
    add_table(
        doc,
        ["Sección del informe", "Imagen/video recomendado", "Por qué aporta valor"],
        [
            ["Introducción", "Captura de login web y login móvil.", "Muestra identidad visual y doble plataforma."],
            ["Análisis y diseño", "Capturas de Figma móvil/web.", "Demuestra planificación previa."],
            ["Base de datos", "Diagrama ER y captura de RDS/Workbench.", "Explica el modelo relacional."],
            ["Backend", "Swagger con endpoints desplegados.", "Evidencia API real y documentada."],
            ["Seguridad", "Flujo JWT dibujado o captura de Authorization Bearer.", "Explica autenticación sin mostrar secretos."],
            ["Móvil", "Video corto escaneando QR y abriendo detalle de máquina.", "Es la funcionalidad más visual del proyecto."],
            ["Web", "Dashboard admin y gestión de usuarios/máquinas.", "Muestra utilidad de escritorio."],
            ["AWS", "Diagrama Amplify -> API Gateway -> EB -> RDS.", "Hace comprensible el despliegue."],
            ["Pruebas", "Captura del smoke test completado.", "Demuestra validación objetiva."],
            ["Entrega", "APK instalada en móvil real.", "Evidencia que no depende solo del emulador."],
        ],
        widths=[1.7, 2.7, 2.3],
    )
    add_note(
        doc,
        "Vídeo de defensa",
        "Recomendación: grabar un vídeo de 60-90 segundos con login web, login móvil, consulta de máquina por QR y vista de Swagger/AWS. Es un recurso muy útil si durante la defensa falla la red o AWS tarda en responder.",
        fill="EEF6FF",
    )

    add_heading(doc, "15. Preguntas de defensa y respuestas preparadas", 1)
    add_table(
        doc,
        ["Pregunta posible", "Respuesta recomendada"],
        [
            ["¿Por qué hay web y móvil?", "Porque tienen usos distintos: web para gestión de escritorio y móvil para interacción física en el gimnasio mediante QR."],
            ["¿Por qué API Gateway si ya existe Elastic Beanstalk?", "Para exponer HTTPS y evitar bloqueo de contenido mixto desde Amplify."],
            ["¿Por qué no usar QR desde web?", "Se puede consultar información, pero escanear QR tiene más sentido en móvil por cámara y contexto físico."],
            ["¿Qué hace que sea seguro?", "BCrypt, JWT, control de roles, CORS configurado y secretos por variables de entorno."],
            ["¿Qué parte fue más compleja?", "Conectar despliegue real web/backend por HTTPS y mantener Android, web y backend usando una misma API."],
            ["¿Qué mejorarías después?", "Dominio propio nativo, HTTPS completo en backend con certificado, recuperación de contraseña, tests automatizados más amplios y paneles analíticos."],
        ],
        widths=[2.4, 4.3],
    )

    add_heading(doc, "16. Anexos e inventario de fuentes", 1)
    add_heading(doc, "16.1 Documentación local", 2)
    md_rows = []
    for md in sorted(DOCS.glob("*.md")):
        md_rows.append([str(md.relative_to(ROOT)).replace("\\", "/"), first_heading(md)])
    add_table(doc, ["Archivo", "Contenido principal"], md_rows[:80], widths=[2.7, 4.0])
    if len(md_rows) > 80:
        add_p(doc, f"Hay {len(md_rows) - 80} documentos Markdown adicionales en la carpeta docs que continúan la bitácora y la documentación técnica.")

    add_heading(doc, "16.2 Fuentes Notion recapituladas", 2)
    add_table(
        doc,
        ["Página / bloque Notion", "Información integrada en el dossier"],
        [
            ["Bitácora del proyecto", "Seguimiento cronológico global, decisiones, incidencias y estado de avance."],
            ["Proyecto intermodular entrega", "Estado global de entrega, bloques completados y pendientes."],
            ["Especificaciones AWS", "RDS, Elastic Beanstalk, Amplify, API Gateway, presupuesto, variables y URLs."],
            ["Actualización frontend móvil Android", "Login, roles, DataStore, QR, rutas, gestión admin, clases, rutinas y asignaciones."],
            ["Fidelidad Figma y navegación interna", "Evolución visual de pantallas y navegación por rol."],
            ["Gestión de usuarios admin", "CRUD de usuarios y validaciones desde app móvil."],
            ["Gestión de clases y rutinas", "Flujo de clases, sesiones, rutinas y asignaciones."],
            ["Asignaciones de rutinas", "Selector con búsqueda y experiencia de asignación entrenador-cliente."],
            ["React, pnpm y frontend web", "Motivo de React, estructura simple, pnpm, variables Vite y despliegue."],
            ["Proxy Amplify API Gateway", "Problema mixed content, solución API Gateway y corrección del /{proxy}."],
            ["Android conectado a API remota", "Migración de 10.0.2.2 a API Gateway para APK instalada en móvil real."],
            ["Smoke test AWS + APK final", "Resultado de validación remota y generación de APK de demo."],
            ["Pulido final y datos demo", "Limpieza de datos de prueba, scripts SQL finales y recursos multimedia."],
        ],
        widths=[2.4, 4.3],
    )
    add_heading(doc, "16.3 Estructura de carpetas", 2)
    add_code(
        doc,
        "backend/        API Spring Boot, seguridad, controladores, entidades, repositorios y scripts de prueba\n"
        "frontend/       Aplicación web React/Vite/pnpm desplegada en Amplify\n"
        "mobile-android/ Aplicación nativa Kotlin/Jetpack Compose y generación de APK\n"
        "database/       Scripts SQL de schema, seed, demo y limpieza\n"
        "docs/           Bitácora técnica, guías, validaciones y este dossier maestro\n"
        ".github/        Workflows de CI/CD y compilación",
    )
    add_heading(doc, "16.4 Checklist final antes de entregar", 2)
    add_bullets(
        doc,
        [
            "Confirmar que Amplify carga y permite login con los tres roles.",
            "Confirmar que Android entra desde móvil real usando API Gateway.",
            "Guardar captura del smoke test remoto exitoso.",
            "Exportar o capturar Swagger y arquitectura AWS.",
            "Preparar vídeo de QR de máquina/sesión si es posible.",
            "Revisar que no se suben node_modules, dist, build ni secretos al repositorio.",
            "Añadir en la memoria final capturas de Figma, ER, API, web, móvil y AWS.",
        ],
    )

    DOCS.mkdir(exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_doc()
